import io
import os
import hashlib
import shutil
import time
from typing import List, Optional, Dict, Any, BinaryIO, Tuple, Union
from fastapi import UploadFile, BackgroundTasks
from sqlalchemy import select, String
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.features.knowledge.models.document import Document, DocumentStatus, DocumentChunk
from app.features.knowledge.models.knowledge import KnowledgeBase
from app.features.knowledge.schemas.request.document import DocumentUpdate
from app.features.knowledge.schemas.request.knowledge import SearchQuery, GetVectorQuery
from app.features.knowledge.service.knowledge import KnowledgeBaseService
from app.features.knowledge.core.config import get_document_settings
from app.features.knowledge.core.response import (
    ServiceResponse,
    ValidationError,
    NotFoundError, ServiceError
)
from app.features.knowledge.core.decorators import (
    handle_service_errors,
    retry
)
from app.features.knowledge.core.logging import (
    service_logger,
    audit_logger,
    metrics_logger
)

class DocumentService:
    """Service for managing documents in knowledge bases"""

    def __init__(
        self,
        db: Session,
        kb_service: KnowledgeBaseService
    ):
        self.db = db
        self.kb_service = kb_service
        self.settings = get_document_settings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.CHUNK_SIZE,
            chunk_overlap=self.settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    @retry(max_attempts=3)
    async def _extract_text(self, file: BinaryIO, file_type: str) -> str:
        """Extract text from a file"""
        try:
            if file_type == "pdf":
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
                return text
            
            elif file_type == "docx":
                doc = docx.Document(file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n\n"
                return text
            
            elif file_type == "txt":
                text = ""
                text = file.read().decode("utf-8")
                return text
            
            else:
                raise ValidationError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            service_logger.error(
                f"Text extraction failed for file type {file_type}",
                "text_extraction_failed",
                error=str(e)
            )
            raise

    async def _split_text(self, text: str, chunk_size: int, chunk_overlap: int, separators: List[str]) -> List[str]:
        """Split text into chunks"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=separators
        )
        chunks = text_splitter.split_text(text)
        if len(chunks) > self.settings.MAX_CHUNKS_PER_DOC:
            raise ValidationError(
                f"Document too large. Maximum chunks allowed: {self.settings.MAX_CHUNKS_PER_DOC}"
            )
        return chunks

    @retry(max_attempts=3)
    async def _store_file(
        self,
        file: UploadFile,
        kb_id: int,
        user_id: int
    ) -> tuple[str, str]:
        """Store uploaded file"""
        try:
            # Validate file size
            file_content = await file.read()
            if len(file_content) > self.settings.MAX_FILE_SIZE:
                raise ValidationError(
                    f"File too large. Maximum size allowed: {self.settings.MAX_FILE_SIZE} bytes"
                )

            # Generate unique filename
            file_hash = hashlib.md5(file_content).hexdigest()

            filepath = await self.kb_service.storage.store_file(
                file_content, kb_id, user_id, file.filename)
            
            return filepath, file_hash
            
        except Exception as e:
            service_logger.error(
                "File storage failed",
                "file_storage_failed",
                error=str(e),
                kb_id=kb_id,
                user_id=user_id
            )
            raise

    async def _process_document(
        self,
        document: Document,
        chunks: List[str],
        kb: KnowledgeBase,
        metadata: Optional[Dict[str, Any]],
        user_id: int
    ):
        """Process document asynchronously to generate embeddings and store vectors"""
        try:
            # Generate embeddings
            embeddings = await self.kb_service.embedding.get_embeddings(chunks)
            chunk_metadata = [{
                "document_id": document.id,
                "chunk_index": i,
                "text": chunk,
                **(metadata or {})
            } for i, chunk in enumerate(chunks)]
            
            # Insert vectors into vector database
            success = await self.kb_service.vector_db.insert_vectors(
                collection_name=kb.collection_name,
                vectors=embeddings,
                metadata=chunk_metadata
            )
            
            if not success:
                raise Exception("Failed to store vectors")
            
            # Update document status
            document.status = DocumentStatus.COMPLETED.value
            document.chunk_count = len(chunks)
            self.db.commit()
            
            # Update knowledge base statistics
            kb.document_count += 1
            kb.total_chunks += len(chunks)
            self.db.commit()
            
            audit_logger.log_access(
                user_id,
                "document",
                str(document.id),
                "process",
                "success"
            )
            
        except Exception as e:
            # Update document status on failure
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)
            self.db.commit()
            service_logger.error(
                "Document processing failed",
                "document_processing_failed",
                error=str(e),
                document_id=document.id
            )

    @handle_service_errors
    async def upload_document(
        self,
        kb_id: int,
        file: UploadFile,
        user_id: int,
        background_tasks: BackgroundTasks,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_rule: Optional[Dict[str, Any]] = None
    ) -> ServiceResponse[Document]:
        """Upload and process a document"""
        start_time = metrics_logger.time()
        
        try:
            # Validate file type
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext[1:] not in self.settings.SUPPORTED_TYPES:
                raise ValidationError(
                    f"Unsupported file type. Supported types: {', '.join(self.settings.SUPPORTED_TYPES)}"
                )
            
            # Get knowledge base
            kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
            if not kb:
                raise NotFoundError("Knowledge base not found")
            exist_document = self.db.query(Document).filter(
                Document.kb_id == kb_id,
                Document.file_name == file.filename).first()
            if exist_document:
                raise ValidationError(f"Document {file.filename} in Knowledge {kb.name} already exists")

            try:
                # Store file
                filepath, file_hash = await self._store_file(file, kb_id, user_id)
                # Create document record with processing status
                # Split text into chunks
                chunk_size = chunk_rule.get("chunk_size",
                                            self.settings.CHUNK_SIZE) if chunk_rule else self.settings.CHUNK_SIZE
                chunk_overlap = chunk_rule.get("chunk_overlap",
                                               self.settings.CHUNK_OVERLAP) if chunk_rule else self.settings.CHUNK_OVERLAP
                separators = chunk_rule.get("separators", ["\n\n", "\n", " ", ""]) if chunk_rule else ["\n\n", "\n",
                                                                                                       " ", ""]
                document = Document(
                    kb_id=kb_id,
                    file_name=file.filename,
                    file_path=filepath,
                    file_type=file_ext[1:],
                    file_size=file.size,
                    file_hash=file_hash,
                    status=DocumentStatus.PROCESSING.value,
                    meta_info=metadata,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    separators=separators,
                    embedding_model=kb.model
                )
                self.db.add(document)
                self.db.commit()
                self.db.refresh(document)
                
                # Extract text
                # with open(filepath, "rb") as f:
                #     text = await self._extract_text(f, file_ext[1:])
                text = await self._extract_text(file.file, file_ext[1:])
                
                # Split text into chunks
                chunks = await self._split_text(
                    text,
                    document.chunk_size,
                    document.chunk_overlap,
                    document.separators
                )
                
                # Store chunks in the database
                for i, chunk in enumerate(chunks):
                    document_chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk,
                        token_count=len(chunk.split())
                    )
                    self.db.add(document_chunk)
                
                self.db.commit()

                # Add background task for processing
                background_tasks.add_task(
                    self._process_document,
                    document,
                    chunks,
                    kb,
                    metadata,
                    user_id
                )

                audit_logger.log_access(
                    user_id,
                    "document",
                    str(document.id),
                    "upload",
                    "success",
                    details={"kb_id": kb_id}
                )
                
                metrics_logger.log_operation(
                    "document_upload",
                    metrics_logger.time() - start_time,
                    True,
                    {"document_id": document.id, "chunks": len(chunks)}
                )
                
                return ServiceResponse.ok(document.to_dict())
                
            except Exception as e:
                # Update document status on failure
                document.status = DocumentStatus.FAILED.value
                document.error_message = str(e)
                self.db.commit()
                
                # Clean up file if stored
                if document.file_path and os.path.exists(document.file_path):
                    try:
                        os.remove(document.file_path)
                    except:
                        pass
                
                metrics_logger.log_operation(
                    "document_upload",
                    metrics_logger.time() - start_time,
                    False,
                    {"error": str(e)}
                )
                raise
                
        except Exception as e:
            self.db.rollback()
            raise

    async def batch_upload_documents(
        self,
        kb_id: int,
        files: List[UploadFile],
        user_id: int,
        background_tasks: BackgroundTasks,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_rule: Optional[Dict[str, Any]] = None
    ) -> ServiceResponse[Document]:
        """Batch upload documents"""
        doc_list = []
        for file in files:
            result = await self.upload_document(kb_id, file, user_id, background_tasks, metadata, chunk_rule)
            doc_list.append(result.data)
        return ServiceResponse.ok(doc_list)

    # @handle_service_errors
    async def get_document(
        self,
        doc_id: int,
        user_id: int
    ) -> Document:
        """Get a document by ID"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
            
        # Check access through knowledge base
        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
            
        audit_logger.log_access(
            user_id,
            "document",
            str(doc_id),
            "read",
            "success"
        )
        
        return document

    @handle_service_errors
    async def delete_document(
        self,
        doc_id: int,
        user_id: int
    ) -> ServiceResponse[None]:
        """Delete a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
            
        # Check access through knowledge base
        kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == document.kb_id).first()
        if not kb:
            raise NotFoundError("Knowledge base not found")
        
        try:
            # Delete vectors from vector database
            success = await self.kb_service.vector_db.delete_vectors(
                collection_name=kb.collection_name,
                metadata_filter={"document_id": doc_id}
            )
            
            if not success:
                print("Failed to delete vectors")
                # raise Exception("Failed to delete vectors")

            # Delete DocumentChunks
            self.db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()

            # Delete file
            # if document.file_path and os.path.exists(document.file_path):
            #     os.remove(document.file_path)
            if document.file_path:
                await self.kb_service.storage.delete_file(document.file_path)
            
            # Update knowledge base statistics
            kb.document_count -= 1
            kb.total_chunks -= document.chunk_count
            
            # Delete document record
            self.db.delete(document)
            self.db.commit()
            
            audit_logger.log_access(
                user_id,
                "document",
                str(doc_id),
                "delete",
                "success"
            )
            
            return ServiceResponse.ok()
            
        except Exception as e:
            self.db.rollback()
            raise

    @handle_service_errors
    async def delete_knowledge_base(
            self,
            kb_id: int,
            user_id: int
    ) -> ServiceResponse:
        """Delete a knowledge"""
        try:
            kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
            if not kb:
                raise NotFoundError("Knowledge base not found")
            # Delete VectorDB
            success = await self.kb_service.vector_db.delete_vectors(kb.collection_name)
            if not success:
                raise Exception("Failed to delete vectors")
            document_all = self.db.query(Document).filter(Document.kb_id == kb_id).all()
            if document_all:
                for document in document_all:
                    # Delete DocumentChunks
                    self.db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()
                    # await self.delete_document(document.id, user_id)
                self.db.query(Document).filter(Document.kb_id == kb_id).delete()
            await self.kb_service.delete_knowledge_base(kb_id, user_id)
            return ServiceResponse.ok()
        except Exception as e:
            self.db.rollback()
            raise

    # @handle_service_errors
    async def list_documents(
        self,
        kb_id: int,
        user_id: int,
        page_num: int = 0,
        page_size: int = 10,
        file_type: Optional[str] = None,
        status: Optional[int] = None,
        search: Optional[str] = None
    ) -> Tuple[List[Document], int]:
        """List documents in a knowledge base"""
        # Check access through knowledge base
        kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
        if not kb:
            raise NotFoundError("Knowledge base not found")
            
        # Build query
        query = self.db.query(Document).filter(Document.kb_id == kb_id)
        
        if file_type:
            query = query.filter(Document.file_type == file_type)
        if status is not None:
            query = query.filter(Document.status == status)

        if search:
            query = query.filter(
                (Document.file_name.ilike(f"%{search}%")) |
                (Document.title.ilike(f"%{search}%"))
            )

        total = query.count()
        skip = (page_num - 1) * page_size
        limit = page_size
        documents = query.offset(skip).limit(limit).all()
        
        return documents, total

    async def update_document(
        self,
        doc_id: int,
        update_data: DocumentUpdate,
        user_id: int
    ):
        """Update a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")

        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
        chunk_rule = update_data.chunk_rule
        if chunk_rule:
            document.chunk_size = chunk_rule.get("chunk_size", document.chunk_size)
            document.chunk_overlap = chunk_rule.get("chunk_overlap", document.chunk_overlap)
            document.separators = chunk_rule.get("separators", document.separators)

        for field, value in update_data.dict(exclude_unset=True).items():
            setattr(document, field, value)
        
        # 提交更改
        self.db.commit()
        self.db.refresh(document)
        
        # 记录审计日志
        audit_logger.log_access(
            user_id,
            "document",
            str(doc_id),
            "update",
            "success"
        )
        
        return document

    # @handle_service_errors
    async def list_chunks(
            self,
            doc_id: int,
            user_id: int,
            page_num: int = 1,
            page_size: int = 10,
            search: Optional[str] = None
    ) -> Tuple[List[DocumentChunk], int]:
        """List chunks of a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
        
        # Check access through knowledge base
        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
        
        query = self.db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id)
        if search:
            query = query.filter(
                (DocumentChunk.content.ilike(f"%{search}%")) |
                (DocumentChunk.chunk_meta_info.cast(String).ilike(f"%{search}%"))
            )

        total = query.count()
        skip = (page_num - 1) * page_size
        limit = page_size
        chunks = query.offset(skip).limit(limit).all()
        
        return chunks, total

    async def get_chunk(
            self,
            chunk_id: int,
            user_id: int
    ) -> DocumentChunk:
        """Get a document chunk by ID"""
        chunk = self.db.query(DocumentChunk).filter(DocumentChunk.id == chunk_id).first()
        if not chunk:
            raise NotFoundError(f"Document chunk {chunk_id} not found")

        # Check access through knowledge base
        document = self.db.query(Document).filter(Document.id == chunk.document_id).first()
        if not document:
            raise NotFoundError("Document not found")

        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")

        audit_logger.log_access(
            user_id,
            "document_chunk",
            str(chunk_id),
            "read",
            "success"
        )

        return chunk

    async def update_chunk(
            self,
            chunk_id: int,
            new_content: str,
            user_id: int
    ) -> DocumentChunk:
        """Update a document chunk by ID"""
        try:
            # Get the chunk by ID
            chunk = self.db.query(DocumentChunk).filter(DocumentChunk.id == chunk_id).first()
            if not chunk:
                raise NotFoundError(f"Document chunk {chunk_id} not found")

            # Get the document associated with the chunk
            document = self.db.query(Document).filter(Document.id == chunk.document_id).first()
            if not document:
                raise NotFoundError("Document not found")

            if new_content == chunk.content:
                raise ServiceError("New content is the same as the original content")

            # Check access through knowledge base
            kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
            if not kb_response.success:
                raise ServiceError(f"Knowledge base {document.kb_id} not found")

            # Update the chunk content
            chunk.content = new_content
            chunk.token_count = len(new_content.split())
            self.db.add(chunk)
            self.db.commit()

            # Get the knowledge base
            kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == document.kb_id).first()
            if not kb:
                raise NotFoundError("Knowledge base not found")

            # Generate new embeddings for the updated chunk
            embeddings = await self.kb_service.embedding.get_embeddings([new_content])
            chunk_metadata = [{
                "document_id": document.id,
                "chunk_index": chunk.chunk_index,
                "text": new_content,
                **(chunk.chunk_meta_info or {})
            }]

            # Delete the existing vector from the vector database
            success = await self.kb_service.vector_db.delete_vectors(
                collection_name=kb.collection_name,
                metadata_filter={"document_id": document.id, "chunk_index": chunk.chunk_index}
            )

            if not success:
                raise Exception("Failed to delete existing vector")

            # Insert the new vector into the vector database
            success = await self.kb_service.vector_db.insert_vectors(
                collection_name=kb.collection_name,
                vectors=embeddings,
                metadata=chunk_metadata
            )

            if not success:
                raise Exception("Failed to store new vector")

            audit_logger.log_access(
                user_id,
                "document_chunk",
                str(chunk_id),
                "update",
                "success"
            )

            return chunk
        except Exception as e:
            self.db.rollback()
            raise ServiceError(f"Failed to update document chunk: {e}")

    @handle_service_errors
    async def clone_knowledge_base(
            self,
            kb_id: int,
            name: str,
            user_id: int,
            background_tasks: BackgroundTasks
    ) -> ServiceResponse[KnowledgeBase]:
        """Clone a knowledge base"""
        new_kb, original_kb = await self.kb_service.clone_knowledge_base(kb_id, name, user_id)

        # create new documents folder
        upload_dir = os.path.join(
            self.settings.UPLOAD_DIR,
            str(user_id),
            str(new_kb.id)
        )
        os.makedirs(upload_dir, exist_ok=True)

        # copy documents and chunks
        original_docs = self.db.query(Document).filter(Document.kb_id == kb_id).all()
        chunk_metadata = []
        for doc in original_docs:
            file_path = doc.file_path
            new_file_path = os.path.join(upload_dir, doc.file_name)
            shutil.copy(file_path, new_file_path)
            new_doc = Document(
                kb_id=new_kb.id,
                file_name=doc.file_name,
                file_path=new_file_path,
                file_type=doc.file_type,
                file_size=doc.file_size,
                file_hash=doc.file_hash,
                status=doc.status,
                meta_info=doc.meta_info
            )
            self.db.add(new_doc)
            self.db.commit()
            self.db.refresh(new_doc)

            doc_chunks = self.db.query(DocumentChunk).filter(DocumentChunk.document_id == doc.id).all()
            chunks = []
            for chunk in doc_chunks:
                new_chunk = DocumentChunk(
                    document_id=new_doc.id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    token_count=chunk.token_count
                )
                chunks.append(chunk.content)
                self.db.add(new_chunk)
            self.db.commit()

            # copy vectors
            # background_tasks.add_task(
            #     self._process_document,
            #     new_doc,
            #     chunks,
            #     new_kb,
            #     doc.meta_info,
            #     user_id
            # )
            chunk_metadata = [{
                "document_id": new_doc.id,
                "chunk_index": chunk.chunk_index,
                "text": chunk.content,
                **(chunk.chunk_meta_info or {})
            } for chunk in doc_chunks]
            vectors = await self.kb_service.vector_db.get_vectors(
                collection_name=original_kb.collection_name,
                metadata_filter={
                    "document_id": doc.id
                }
            )
            # Insert vectors into vector database
            success = await self.kb_service.vector_db.insert_vectors(
                collection_name=new_kb.collection_name,
                vectors=vectors,
                metadata=chunk_metadata
            )

        return ServiceResponse(
            success=True,
            code=201,
            message="Cloned",
            data=new_kb.to_dict()
        )

    @handle_service_errors
    async def reprocess_document(
            self,
            doc_id: int,
            user_id: int,
            background_tasks: BackgroundTasks
    ) -> ServiceResponse[Document]:
        """Reprocess a document"""
        start_time = metrics_logger.time()

        try:
            # Get document by ID
            document = self.db.query(Document).filter(Document.id == doc_id).first()
            if not document:
                raise NotFoundError(f"Document {doc_id} not found")

            # Check access through knowledge base
            kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
            if not kb_response.success:
                raise ServiceError(f"Knowledge base {document.kb_id} not found")

            # Get knowledge base
            kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == document.kb_id).first()
            if not kb:
                raise NotFoundError("Knowledge base not found")

            # Update document status to PROCESSING
            document.status = DocumentStatus.PROCESSING.value
            document.error_message = None
            self.db.commit()

            # Read text from the document file using the storage provider
            file_content = await self.kb_service.storage.read_file(document.file_path)
            # Convert bytes to file-like object
            file_like_object = io.BytesIO(file_content)
            # Extract text from the document file
            text = await self._extract_text(file_like_object, document.file_type)

            # Split text into chunks
            chunks = await self._split_text(text, document.chunk_size, document.chunk_overlap, document.separators)

            # delete old chunks
            self.db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document.id).delete(synchronize_session=False)
            # Store chunks in the database
            for i, chunk in enumerate(chunks):
                document_chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk,
                    token_count=len(chunk.split())
                )
                self.db.add(document_chunk)
            self.db.commit()

            # Add background task for processing
            background_tasks.add_task(
                self._process_document,
                document,
                chunks,
                kb,
                document.meta_info,
                user_id
            )

            audit_logger.log_access(
                user_id,
                "document",
                str(document.id),
                "reprocess",
                "success"
            )

            metrics_logger.log_operation(
                "document_reprocess",
                metrics_logger.time() - start_time,
                True,
                {"document_id": document.id, "chunks": len(chunks)}
            )

            return ServiceResponse.ok(document.to_dict())

        except Exception as e:
            # Update document status to FAILED
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)
            self.db.commit()

            metrics_logger.log_operation(
                "document_reprocess",
                metrics_logger.time() - start_time,
                False,
                {"error": str(e)}
            )
            raise

    async def download_document(
            self,
            doc_id: int,
            user_id: int
    ) -> StreamingResponse:
        """Download original document"""
        start_time = metrics_logger.time()

        try:
            # Get document by ID
            document = self.db.query(Document).filter(Document.id == doc_id).first()
            if not document:
                raise NotFoundError(f"Document {doc_id} not found")

            # Check access through knowledge base
            kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
            if not kb_response.success:
                raise ServiceError(f"Knowledge base {document.kb_id} not found")

            # Check if file exists
            if not document.file_path:
                raise NotFoundError(f"File path for document {doc_id} is not available")

            # Read file content using the storage provider
            file_content = await self.kb_service.storage.read_file(document.file_path)
            # Convert bytes to file-like object
            file_like_object = io.BytesIO(file_content)
            file_ext = document.file_type
            media_type = "application/octet-stream"
            # if file_ext == "pdf":
            #     media_type = "application/pdf"
            # elif file_ext == "docx":
            #     media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            # elif file_ext == "txt":
            #     media_type = "text/plain"
            # Create a StreamingResponse to return the file
            response = StreamingResponse(file_like_object, media_type=media_type)
            response.headers["Content-Disposition"] = f"attachment; filename={document.file_name}"

            audit_logger.log_access(
                user_id,
                "document",
                str(document.id),
                "download",
                "success"
            )

            metrics_logger.log_operation(
                "document_download",
                metrics_logger.time() - start_time,
                True,
                {"document_id": document.id}
            )

            return response

        except Exception as e:
            metrics_logger.log_operation(
                "document_download",
                metrics_logger.time() - start_time,
                False,
                {"error": str(e)}
            )
            raise

    @handle_service_errors(wrap_response=False)
    async def get_vectors(
            self,
            kb_id: int,
            query: GetVectorQuery
    ):
        """Get documents in a knowledge base"""
        kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()

        if not kb:
            raise NotFoundError(f"Knowledge base {kb_id} not found")
        filters = query.metadata_filter
        vectors = await self.kb_service.vector_db.get_vectors(
            collection_name=kb.collection_name,
            metadata_filter=filters
        )
        return vectors
